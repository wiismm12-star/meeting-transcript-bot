from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def write_docx(
    path: Path,
    title: str,
    text: str,
    *,
    meeting_id: str | None = None,
    meeting_date: str | None = None,
    speakers: list[str] | None = None,
) -> None:
    document = Document()
    _configure_document(document)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, "Microsoft JhengHei", 20, bold=True, color="1F4E79")
    title_paragraph.paragraph_format.space_after = Pt(10)

    metadata = document.add_paragraph(style="Meeting Metadata")
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(f"日期：{_format_meeting_date(meeting_date)}")
    if meeting_id:
        metadata.add_run(f"　｜　會議 ID：{meeting_id}")

    if speakers:
        speaker_paragraph = document.add_paragraph(style="Speaker List")
        label = speaker_paragraph.add_run("主講人：")
        _set_run_font(label, "Microsoft JhengHei", 10.5, bold=True, color="1F4E79")
        speaker_paragraph.add_run("、".join(speakers))

    for block in text.split("\n\n"):
        _add_transcript_block(document, block.strip())

    document.save(path)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(7)

    for style_name, size, color in (("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "1F4E79")):
        style = document.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("Meeting Metadata", "Speaker List"):
        style = document.styles.add_style(style_name, 1)
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(89, 89, 89)
        style.paragraph_format.space_after = Pt(5)


def _add_transcript_block(document: Document, block: str) -> None:
    if not block:
        return

    if block.startswith("# "):
        document.add_heading(block[2:].strip(), level=1)
        return
    if block.startswith("## "):
        document.add_heading(block[3:].strip(), level=2)
        return

    speaker_match = re.match(r"^(.{1,80}?)[：:]\s*(.+)$", block, flags=re.DOTALL)
    if speaker_match and not block.startswith("- "):
        paragraph = document.add_paragraph()
        speaker_run = paragraph.add_run(f"{speaker_match.group(1).strip()}：")
        _set_run_font(speaker_run, "Microsoft JhengHei", 11, bold=True, color="1F4E79")
        paragraph.add_run(speaker_match.group(2).strip())
        return

    if block.startswith("- "):
        document.add_paragraph(block[2:].strip(), style="List Bullet")
        return

    document.add_paragraph(block)


def _format_meeting_date(value: str | None) -> str:
    if not value:
        return "未提供"
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%Y-%m-%d")
    except ValueError:
        return value[:10]


def _set_run_font(run, font_name: str, size: float, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
